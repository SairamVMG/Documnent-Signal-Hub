"""
modules/word_parser.py
DOCX introspection, text extraction, block extraction, and field parsing.
Treats Word files as unstructured/semi-structured documents similar to PDF.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

import docx2txt
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# Public parse API
# ─────────────────────────────────────────────────────────────────────────────

def parse_word(
    file_path: str | Path,
    llm_client=None,
) -> dict:
    """
    Parse a .docx file and return a structure similar to PDF parsing output.

    Returns:
        {
            "doc_type": "word_document",
            "doc_label": "Word Document",
            "raw_text": "...",
            "blocks": [...],
            "fields": [...]
        }
    """
    blocks = extract_word_blocks(file_path)
    raw_text = "\n".join(
        b["text"] for b in blocks if b.get("text", "").strip()
    ).strip()

    fields = extract_word_fields_from_blocks(blocks, llm_client=llm_client)

    return {
        "doc_type": "word_document",
        "doc_label": "Word Document",
        "raw_text": raw_text,
        "blocks": blocks,
        "fields": fields,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Basic helpers
# ─────────────────────────────────────────────────────────────────────────────

def get_word_sheet_names(file_path: str | Path) -> list[str]:
    """
    DOCX doesn't have sheets. Return a single logical document tab.
    """
    return ["Document"]


def get_word_dimensions(file_path: str | Path, section_name: str = "Document") -> tuple[int, int]:
    """
    Return pseudo-dimensions for DOCX.
    rows = number of extracted blocks
    cols = 1
    """
    blocks = extract_word_blocks(file_path)
    return len(blocks), 1


def extract_word_text(file_path: str | Path) -> str:
    """
    Fallback plain-text extraction using docx2txt.
    """
    try:
        return (docx2txt.process(str(file_path)) or "").strip()
    except Exception:
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Block extraction (paragraphs + table cells)
# ─────────────────────────────────────────────────────────────────────────────

def extract_word_blocks(file_path: str | Path) -> list[dict]:
    """
    Extract Word document content as ordered blocks for traceability and highlighting.

    Returns a list like:
    [
        {
            "block_id": 1,
            "block_type": "paragraph",
            "text": "Policy Number: CGL-2021-00847",
            "para_index": 5,
            "table_index": None,
            "row_index": None,
            "col_index": None,
        },
        ...
    ]
    """
    file_path = str(file_path)
    doc = Document(file_path)

    blocks: list[dict] = []
    block_id = 1

    # Paragraph blocks
    for p_idx, para in enumerate(doc.paragraphs):
        txt = (para.text or "").strip()
        if txt:
            blocks.append({
                "block_id": block_id,
                "block_type": "paragraph",
                "text": txt,
                "para_index": p_idx,
                "table_index": None,
                "row_index": None,
                "col_index": None,
            })
            block_id += 1

    # Table cell blocks
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_cells = [((cell.text or "").strip()) for cell in row.cells]
            row_text = " | ".join([c for c in row_cells if c])

            if row_text:
                blocks.append({
                    "block_id": block_id,
                    "block_type": "table_row",
                    "text": row_text,
                    "para_index": None,
                    "table_index": t_idx,
                    "row_index": r_idx,
                    "col_index": None,
                })
                block_id += 1

            for c_idx, cell in enumerate(row.cells):
                txt = (cell.text or "").strip()
                if txt:
                    blocks.append({
                        "block_id": block_id,
                        "block_type": "table_cell",
                        "text": txt,
                        "para_index": None,
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "col_index": c_idx,
                    })
                    block_id += 1

    return blocks


# ─────────────────────────────────────────────────────────────────────────────
# Field extraction
# ─────────────────────────────────────────────────────────────────────────────

_FIELD_LABEL_PATTERNS = [
    "claim number", "claim no", "claim #", "claim id",
    "policy number", "policy no", "policy #",
    "insured", "carrier", "loss date", "date of loss",
    "date reported", "status", "claimant", "claimant name",
    "description of loss", "total paid", "reserve", "total incurred",
    "effective date", "expiration date", "lob", "line of business",
]

_CANONICAL_MAP = {
    "claim number": "Claim Number",
    "claim no": "Claim Number",
    "claim #": "Claim Number",
    "claim id": "Claim Number",
    "policy number": "Policy Number",
    "policy no": "Policy Number",
    "policy #": "Policy Number",
    "insured": "Insured",
    "carrier": "Carrier",
    "loss date": "Loss Date",
    "date of loss": "Loss Date",
    "date reported": "Date Reported",
    "status": "Status",
    "claimant": "Claimant Name",
    "claimant name": "Claimant Name",
    "description of loss": "Description of Loss",
    "total paid": "Total Paid",
    "reserve": "Reserve",
    "total incurred": "Total Incurred",
    "effective date": "Effective Date",
    "expiration date": "Expiration Date",
    "lob": "Line of Business",
    "line of business": "Line of Business",
}


def _canonical_field_name(label: str) -> str:
    k = re.sub(r"\s+", " ", label.strip().lower())
    return _CANONICAL_MAP.get(k, label.strip().title())

def _looks_like_label(text: str) -> bool:
    if not text:
        return False
    t = text.strip().lower()
    return (
        len(t) <= 60
        and any(k in t for k in [
            "policy", "claim", "insured", "carrier", "date", "status",
            "premium", "limit", "deductible", "period", "name", "number",
            "less", "add", "tax"
        ])
    )
def extract_word_fields_from_blocks(blocks: list[dict], llm_client=None) -> list[dict]:
    """
    Extract fields from Word blocks using:
      1. label: value patterns
      2. two-column table rows
    """
    fields: list[dict] = []
    seen_names: set[str] = set()

    # Pass 1: "Label: Value"
    for b in blocks:
        text = b.get("text", "").strip()
        if not text:
            continue

        m = re.match(r"^\s*([A-Za-z0-9 #/_\-\(\)\.]{2,60})\s*:\s*(.+?)\s*$", text)
        if m:
            raw_label = m.group(1).strip()
            raw_value = m.group(2).strip()

            if len(raw_value) >= 1:
                field_name = _canonical_field_name(raw_label)
                if field_name not in seen_names:
                    fields.append({
                        "field_name": field_name,
                        "value": raw_value,
                        "confidence": 0.95,
                        "source_block": b["block_id"],
                        "source_text": text,
                        "source_para": b.get("para_index"),
                        "source_table": b.get("table_index"),
                        "source_row": b.get("row_index"),
                        "source_col": b.get("col_index"),
                    })
                    seen_names.add(field_name)

    # Pass 2: table rows like "Policy Number | CGL-2021-00847"
    for b in blocks:
        if b.get("block_type") != "table_row":
            continue

        text = b.get("text", "").strip()
        parts = [p.strip() for p in text.split("|") if p.strip()]
        if len(parts) >= 2:
            raw_label = parts[0]
            raw_value = parts[1]

            if raw_label.lower() in _FIELD_LABEL_PATTERNS and raw_value:
                field_name = _canonical_field_name(raw_label)
                if field_name not in seen_names:
                    fields.append({
                        "field_name": field_name,
                        "value": raw_value,
                        "confidence": 0.92,
                        "source_block": b["block_id"],
                        "source_text": text,
                        "source_para": b.get("para_index"),
                        "source_table": b.get("table_index"),
                        "source_row": b.get("row_index"),
                        "source_col": b.get("col_index"),
                    })
                    seen_names.add(field_name)
        # Pass 2B: adjacent table-cell label/value pairing
    table_cells = [b for b in blocks if b.get("block_type") == "table_cell"]

    for i in range(len(table_cells) - 1):
        left = table_cells[i]
        right = table_cells[i + 1]

        # same table and same row
        if (
            left.get("table_index") == right.get("table_index")
            and left.get("row_index") == right.get("row_index")
            and left.get("col_index") is not None
            and right.get("col_index") is not None
            and right.get("col_index") == left.get("col_index") + 1
        ):
            raw_label = (left.get("text") or "").strip()
            raw_value = (right.get("text") or "").strip()

            if _looks_like_label(raw_label) and raw_value and len(raw_value) <= 200:
                field_name = _canonical_field_name(raw_label)
                if field_name not in seen_names:
                    fields.append({
                        "field_name": field_name,
                        "value": raw_value,
                        "confidence": 0.93,
                        "source_block": right["block_id"],
                        "source_text": f"{raw_label}: {raw_value}",
                        "source_para": right.get("para_index"),
                        "source_table": right.get("table_index"),
                        "source_row": right.get("row_index"),
                        "source_col": right.get("col_index"),
                    })
                    seen_names.add(field_name)

    # Pass 3: fallback regex scan across all text if nothing found
    if not fields:
        joined = "\n".join([b.get("text", "") for b in blocks])

        fallback_patterns = [
            (r"\bPolicy\s*(?:Number|No|#)\s*[:\-]?\s*([A-Z0-9\-_\/]+)", "Policy Number"),
            (r"\bClaim\s*(?:Number|No|#|ID)\s*[:\-]?\s*([A-Z0-9\-_\/]+)", "Claim Number"),
            (r"\bInsured\s*[:\-]?\s*([^\n\r]+)", "Insured"),
            (r"\bCarrier\s*[:\-]?\s*([^\n\r]+)", "Carrier"),
            (r"\bLoss\s*Date\s*[:\-]?\s*([^\n\r]+)", "Loss Date"),
            (r"\bEffective\s*Date\s*[:\-]?\s*([^\n\r]+)", "Effective Date"),
        ]

        for pat, name in fallback_patterns:
            m = re.search(pat, joined, re.IGNORECASE)
            if m:
                val = m.group(1).strip()
                fields.append({
                    "field_name": name,
                    "value": val,
                    "confidence": 0.80,
                    "source_block": None,
                    "source_text": val,
                    "source_para": None,
                    "source_table": None,
                    "source_row": None,
                    "source_col": None,
                })

    return fields